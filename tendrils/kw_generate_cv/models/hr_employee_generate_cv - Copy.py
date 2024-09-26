from datetime import datetime, date
import html2text
from docx import Document
from htmldocx import HtmlToDocx
from odoo import models, fields, api, tools, http


class GenerateCV(models.Model):
    _inherit = 'hr.employee'

    def view_cv_button(self):
        view_id = self.env.ref('kw_generate_cv.hr_employee_generate_cv_form_view').id
        return {
            'name': 'Generate CV',
            'type': 'ir.actions.act_window',
            'view_type': 'form',
            'view_mode': 'form',
            'views': [(view_id, 'form')],
            'res_model': 'hr.employee',
            'view_id': view_id,
            'target': 'current',
            'context': {'create': False, 'edit': False, 'delete': False, "toolbar": False},
            'res_id': self.id,
        }

    def button_cv_download(self):
        # print("button_generate_cv_docx ???????????????? ")
        document = Document()
        document.add_heading("Aspernatur dolor excepturi temporibus ut vero.", 0)
        # new_parser = HtmlToDocx()
        # print("button_generate_cv_docx ???????????????? ", document, new_parser)
        # # do stuff to document
        #
        # html = """<h2>Aspernatur dolor excepturi temporibus ut vero.</h2>
        # <p>Aspernatur dolor excepturi temporibus ut vero. Alias at beatae
        #  dolorum impedit nemo nisi non numquam, omnis quidem quis repellat sed similique unde!</p>"""
        # new_parser.add_html_to_document(html, document)

        # do more stuff to document
        document.save('abc.docx')
        docx = document
        docxhttpheaders = [
            ('Content-Type', 'application/vnd.openxmlformats-'
                             'officedocument.wordprocessingml.document'),
            ('Content-Length', len(docx)),
            (
                'Content-Disposition',
                content_disposition('sample cv.docx')
            )
        ]
        return http.request.make_response(docx, headers=docxhttpheaders)


class GenerateCVReport(models.AbstractModel):
    _name = 'report.kw_generate_cv.report_generate_cv_docx'
    _inherit = 'report.report_docx.abstract'

    def generate_docx_report(self, document, data, record):
        # print("document >>>>> ", document)
        # print("data >>>>> ", data)
        # print("agreement_closure >>>>> ", agreement_closure)
        # if agreement_closure.template_ids:
        #     for section in agreement_closure.template_ids:
        # document.add_section("""<h2>Aspernatur dolor excepturi temporibus ut vero.</h2>
        # <p>Aspernatur dolor excepturi temporibus ut vero. Alias at beatae
        #  dolorum impedit nemo nisi non numquam, omnis quidem quis repellat sed similique unde!</p>""")
        # document.add_heading("Aspernatur dolor excepturi temporibus ut vero.", 0)
        # document.add_paragraph(html2text.html2text("""Aspernatur dolor excepturi temporibus ut vero. Alias at beatae
        # dolorum impedit nemo nisi non numquam, omnis quidem quis repellat sed similique unde!"""))

        # document.add_heading("Aspernatur dolor excepturi temporibus ut vero.", 0)
        # document.add_paragraph(html2text.html2text("""Aspernatur dolor excepturi temporibus ut vero. Alias at beatae
        # dolorum impedit nemo nisi non numquam, omnis quidem quis repellat sed similique unde!"""))
        new_parser = HtmlToDocx()
        new_parser.table_style = 'TableGrid'

        html = f"""<table border="1" style="width: 100%; border:1px solid #000;" >
       
    <tr>
        <td style="border: 1px solid #000000; width:5%">1</td>
        <td style="border: 1px solid #000000; width:30%">Proposed Position</td>
        <td style="border: 1px solid #000000; width:65%"><strong>{record.job_id.name}</strong> </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">2</td>
        <td style="border: 1px solid #000000; width:30%">Name of Firm</td>
        <td style="border: 1px solid #000000; width:65%">CSM Technologies</td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">3</td>
        <td style="border: 1px solid #000000; width:30%">Name of Staff</td>
        <td style="border: 1px solid #000000; width:65%"><b>{record.name}</b></td>
    </tr>
    # <tr>
    #     <td style="border: 1px solid #000000; width:5%">4</td>
    #     <td style="border: 1px solid #000000; width:30%">Date of Birth</td>
    #     <td style="border: 1px solid #000000; width:65%">{datetime.strptime(str(record.birthday), "%Y-%m-%d").strftime("%d-%b-%Y")}</td>
    # </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">5</td>
        <td style="border: 1px solid #000000; width:30%">Nationality</td>
        <td style="border: 1px solid #000000; width:65%">ssss</td>
    </tr>"""
        # new_parser.add_html_to_document(html, document)
        html += """
    <tr>
        <td style="border: 1px solid #000000; width:5%">6</td>
        <td style="border: 1px solid #000000; width:30%">Education</td>
        <td style="border: 1px solid #000000; width:65%">
        <table border="1" >
            <tr>
                <th >Course Name</th>
                <th >Stream</th>
                <th >University</th>
                <th >Passing Year</th>
            </tr>
            <tr >
                <td >Matriculation</td>
                <td >10TH</td>
                <td >BSE</td>
                <td >2008</td>
            </tr>
            <tr >
                <td >Intermediate</td>
                <td >+2 SCIENCE</td>
                <td >CBSE</td>
                <td >2010</td>
            </tr>
            <tr >
                <td >Graduation</td>
                <td >BCA</td>
                <td >BPUT</td>
                <td >2013</td>
            </tr>
            <tr >
                <td>Post Graduation</td>
                <td>MCA</td>
                <td>BPUT</td>
                <td>2016</td>
            </tr>
            </table>
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">7</td>
        <td style="border: 1px solid #000000; width:30%">Membership of Professional Organization</td>
        <td style="border: 1px solid #000000; width:65%">
            <table border="1">
                <tr>
                    <th>Association Name</th>
                    <th>Date of Issue</th>
                    <th>Date of Expiry</th>
                    <th>Renewal Applied</th>
                </tr>
                <tr><td colspan="4">&nbsp;</td></tr>
                <tr><td colspan="4">&nbsp;</td></tr>
                <tr><td colspan="4">&nbsp;</td></tr>
                <tr><td colspan="4">&nbsp;</td></tr>
            </table>
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">9</td>
        <td style="border: 1px solid #000000; width:30%">Countries Of Work Experience</td>
        <td style="border: 1px solid #000000; width:65%">
       
            <table border="1">
            <tr>
                <th>Country Name</th>
                <th>Country Code</th>
                <th>Active</th>
            </tr>
                <tr><td colspan="3">&nbsp;</td></tr>
                <tr><td colspan="3">&nbsp;</td></tr>
                <tr><td colspan="3">&nbsp;</td></tr>
                <tr><td colspan="3">&nbsp;</td></tr>
            </table>
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">10</td>
        <td style="border: 1px solid #000000; width:30%">Languages</td>
        <td style="border: 1px solid #000000; width:65%">
            <table border="1">
                <tr>
                    <th>Language</th>
                    <th>Reading</th>
                    <th>Writing</th>
                    <th>Speaking</th>
                    <th>Understanding</th>
                </tr>
                <tr><td colspan="5">&nbsp;</td></tr>
                <tr><td colspan="5">&nbsp;</td></tr>
                <tr><td colspan="5">&nbsp;</td></tr>
                <tr><td colspan="5">&nbsp;</td></tr>
            </table>
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%">11</td>
        <td style="border: 1px solid #000000; width:30%">Employeement Record</td>
        <td style="border: 1px solid #000000; width:65%">
            <table border="1">
                <tr>
                    <th>Country Name</th>
                    <th>Previous Organization Name</th>
                    <th>Job Profile </th>
                    <th>Organization Type</th>
                    <th>Industry Type</th>
                    <th>Effective From</th>
                    <th>Effective To</th>
                </tr>
                <tr>
                    <td>Afghanistan</td>
                    <td>asdas</td>
                    <td>asd</td>
                    <td>Defence</td>
                    <td >Banking</td>
                    <td>01-Jan-2021</td>
                    <td>31-Jul-2022</td>
                </tr>
                <tr><td colspan="7">&nbsp;</td></tr>
                <tr><td colspan="7">&nbsp;</td></tr>
                <tr><td colspan="7">&nbsp;</td></tr>
            </table>
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%;">13</td>
        <td style="border: 1px solid #000000; width:30%;">
            Work Undertaken that Illustrates Capability to Assigned Handle the Task Assigned
        </td>
        <td style="border: 1px solid #000000; width:65%;">
       
            <table border="1">
                <tr><th >ID</th></tr>
                <tr><td colspan="1">&nbsp;</td></tr><tr><td colspan="1">&nbsp;</td></tr>
                <tr><td colspan="1">&nbsp;</td></tr><tr><td colspan="1">&nbsp;</td></tr>
            </table>
        
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; width:5%;">14</td>
        <td style="border: 1px solid #000000; width:30%;">Certification</td>
        <td style="border: 1px solid #000000; width:65%;">
            I, the undersigned, certify that to the best of my knowldge and belief,
            this cv correctly describes me, my qualifications and my exprience. I
            understand that any willfull misstatement described herein may lead to my
            disqualification or dissmisal, if engaged.
        
        </td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; padding:5px;">Signature: Asit Mohanty </td>
        <td style="border: 1px solid #000000; padding:5px;">Signature:</td>
        <td style="border: 1px solid #000000; padding:5px;"></td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; padding:5px;">Date: 12-Sep-2022</td>
        <td style="border: 1px solid #000000;">Date: 12-Sep-2022</td>
        <td style="border: 1px solid #000000;"></td>
    </tr>
    <tr>
        <td style="border: 1px solid #000000; padding:5px;">Name of Staff Member:Asit Mohanty</td>
        <td style="border: 1px solid #000000; padding:5px;">Name of Authorized Signatory:</td>
        <td style="border: 1px solid #000000; padding:5px;"></td>
    </tr>
    </table>"""
        # new_parser = HtmlToDocx()
        # new_parser.table_style = 'TableGrid'
        new_parser.add_html_to_document(html, document)
