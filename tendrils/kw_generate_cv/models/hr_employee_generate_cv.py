from datetime import datetime, date
import html2text
from docx import Document
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm, Pt
from htmldocx import HtmlToDocx
from odoo import models, fields, api, tools, http


class GenerateCV(models.Model):
    _inherit = 'hr.employee'

    def view_cv_button(self):
        view_id = self.env.ref('kw_generate_cv.hr_employee_generate_cv_form_view').id
        return {
            'name': 'Generate CV',
            'type': 'ir.actions.act_window',
            'view_type': 'form',
            'view_mode': 'form',
            'views': [(view_id, 'form')],
            'res_model': 'hr.employee',
            'view_id': view_id,
            'target': 'current',
            'context': {'create': False, 'edit': False, 'delete': False, "toolbar": False},
            'res_id': self.id,
        }

    def button_cv_download(self):
        # print("button_generate_cv_docx ???????????????? ")
        document = Document()
        document.add_heading("Aspernatur dolor excepturi temporibus ut vero.", 0)
        # new_parser = HtmlToDocx()
        # print("button_generate_cv_docx ???????????????? ", document, new_parser)
        # # do stuff to document
        #
        # html = """<h2>Aspernatur dolor excepturi temporibus ut vero.</h2>
        # <p>Aspernatur dolor excepturi temporibus ut vero. Alias at beatae
        #  dolorum impedit nemo nisi non numquam, omnis quidem quis repellat sed similique unde!</p>"""
        # new_parser.add_html_to_document(html, document)

        # do more stuff to document
        document.save('abc.docx')
        docx = document
        # print("document >>> ", docx, document)
        docxhttpheaders = [
            ('Content-Type', 'application/vnd.openxmlformats-'
                             'officedocument.wordprocessingml.document'),
            ('Content-Length', len(docx)),
            (
                'Content-Disposition',
                content_disposition('sample cv.docx')
            )
        ]
        return http.request.make_response(docx, headers=docxhttpheaders)


class GenerateCVReport(models.AbstractModel):
    _name = 'report.kw_generate_cv.report_generate_cv_docx'
    _inherit = 'report.report_docx.abstract'

    def generate_docx_report(self, document, data, record):
        # print("document >>>>> ", document)
        # print("data >>>>> ", data)
        # print("agreement_closure >>>>> ", agreement_closure)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        heading = document.add_heading(f"{record.job_id.name}: {record.name}", 1)
        heading.style = 'Title'
        heading.style.font.size = Pt(14)

        p = document.add_paragraph('', 'List Number')
        p.add_run('Proposed Position: ').bold = True
        p.add_run(record.job_id.name)

        p = document.add_paragraph('', 'List Number')
        p.add_run('Name of Firm: ').bold = True
        p.add_run('CSM Technologies')

        p = document.add_paragraph('', 'List Number')
        p.add_run('Name of Staff: ').bold = True
        p.add_run(record.name)

        p = document.add_paragraph('', 'List Number')
        p.add_run('Date of Joining: ').bold = True
        p.add_run(datetime.strptime(str(record.date_of_joining), "%Y-%m-%d").strftime("%d-%b-%Y"))

        p = document.add_paragraph('', 'List Number')
        p.add_run('Nationality: ').bold = True
        p.add_run(record.country_id.name or 'NA')

        new_parser = HtmlToDocx()
        new_parser.table_style = 'TableGrid'
        # new_parser.add_html_to_document(html, document)

        # 6. Qualification
        p = document.add_paragraph('', 'List Number')
        p.add_run('Qualification: ').bold = True
        education_records = record.educational_details_ids.filtered(lambda x: x.course_type == '1')
        if education_records:
            rows_len = len(education_records) + 1
            table = document.add_table(rows=rows_len, cols=3, style='Table Grid')
            table.rows[0].cells[0].paragraphs[0].add_run('College/ University/ Institutions').font.bold = True
            table.rows[0].cells[1].paragraphs[0].add_run('Degrees Obtained').font.bold = True
            table.rows[0].cells[2].paragraphs[0].add_run('Dates Of Obtainment').font.bold = True
            cntr = 1
            for rec in education_records:
                table.rows[cntr].cells[0].text = rec.university_name.name
                table.rows[cntr].cells[1].text = rec.stream_id.name
                table.rows[cntr].cells[2].text = rec.passing_year
                table.rows[cntr].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cntr += 1
        else:
            p.add_run('NA')

        # 7. Certifications
        p = document.add_paragraph('', 'List Number')
        p.add_run('Certifications: ').bold = True
        education_records = record.educational_details_ids.filtered(lambda x: x.course_type != '1')
        if education_records:
            rows_len = len(education_records) + 1
            table = document.add_table(rows=rows_len, cols=3, style='Table Grid')
            table.rows[0].cells[0].paragraphs[0].add_run('Course/ Certifications').font.bold = True
            table.rows[0].cells[1].paragraphs[0].add_run('Institute').font.bold = True
            table.rows[0].cells[2].paragraphs[0].add_run('Dates Of Obtainment').font.bold = True
            cntr = 1
            for rec in education_records:
                table.rows[cntr].cells[0].text = rec.stream_id.name
                table.rows[cntr].cells[1].text = rec.university_name.name
                table.rows[cntr].cells[2].text = rec.passing_year
                table.rows[cntr].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cntr += 1
        else:
            p.add_run('NA')

        # 8. Countries of Work Experience
        p = document.add_paragraph('', 'List Number')
        p.top_margin = Pt(10)
        p.add_run('Work Experience: ').bold = True
        p.add_run(', '.join([rec.name for rec in record.work_experience_ids]) if record.work_experience_ids else 'NA')

        # 9. language
        p = document.add_paragraph('', 'List Number')
        p.add_run('Languages: ').bold = True
        if record.known_language_ids:
            html = "<table> <tr> <th>Language</th> <th>Speaking</th> <th>Reading</th> <th>Writing</th> </tr>"
            for rec in record.known_language_ids:
                html += f"""<tr>
                    <th>{rec.language_id.name}</th>
                    <td>{rec.speaking_status.capitalize()}</td>
                    <td>{rec.reading_status.capitalize()}</td>
                    <td>{rec.writing_status.capitalize()}</td>
                </tr>"""
            html += "</table>"
            new_parser.add_html_to_document(html, document)
        else:
            p.add_run('NA')

        # 10. Employment Record
        p = document.add_paragraph('', 'List Number')
        p.add_run('Employment Record: ').bold = True

        p = document.add_paragraph('', 'List Number 2')
        p.add_run('From: ').bold = True
        p.add_run(datetime.strptime(str(record.date_of_joining), "%Y-%m-%d").strftime("%b %Y"))
        p.add_run(' to till date \n')
        p.add_run('Employer: ').bold = True
        p.add_run('CSM Technologies \n')
        p.add_run('Positions held: ').bold = True
        p.add_run(record.job_id.name)

        if record.work_experience_ids:
            for rec in record.work_experience_ids:
                p = document.add_paragraph('', 'List Number 2')
                p.add_run('From: ').bold = True
                p.add_run(datetime.strptime(str(rec.effective_from), "%Y-%m-%d").strftime("%b %Y"))
                p.add_run(' to ')
                p.add_run(datetime.strptime(str(rec.effective_to), "%Y-%m-%d").strftime("%b %Y"))
                p.add_run('\nEmployer: ').bold = True
                p.add_run(f'{rec.name} \n')
                p.add_run('Positions held: ').bold = True
                p.add_run(rec.designation_name)

        # 11. Work Undertaken
        p = document.add_paragraph('', 'List Number')
        p.add_run('Work Undertaken that Best Illustrates Capability to Assigned Handle the Tasks Assigned: ').bold = True

        if record.cv_info_details_ids:
            cntr = 1
            for rec in record.cv_info_details_ids:
                table = document.add_table(rows=9, cols=2, style='Table Grid')
                headers = [f'Project {cntr}', 'Year', 'Location', 'Client', 'Main Project Features', '', 'Position Held',
                           'Activities Performed', '']
                cntr += 1
                # table.style.border = Pt(1)
                for key, row in enumerate(table.rows):
                    # for idx, width in enumerate(widths):
                    if key in [4, 5, 7, 8]:
                        row.cells[0].merge(row.cells[-1])
                        if key in [4, 7]:
                            row.cells[0].paragraphs[0].add_run(headers[key]).font.bold = True
                        else:
                            if key == 5:
                                # row.cells[0].text = rec.project_feature
                                new_parser.add_html_to_document(rec.project_feature, row.cells[0])
                            elif key == 8:
                                # row.cells[0].text = rec.responsibility_activity
                                new_parser.add_html_to_document(rec.responsibility_activity, row.cells[0])
                    else:
                        row.cells[0].width = Cm(4)
                        # row.cells[0].text = headers[key]
                        row.cells[0].paragraphs[0].add_run(headers[key]).font.bold = True
                        if key == 0:
                            row.cells[1].text = rec.compute_project if rec.compute_project else 'NA'
                        elif key == 1:
                            if rec.start_month_year and rec.end_month_year:
                                row.cells[1].text = f"{datetime.strptime(str(rec.start_month_year), '%Y-%m-%d').strftime('%b %Y')} to {datetime.strptime(str(rec.end_month_year), '%Y-%m-%d').strftime('%b %Y')}"
                            else:
                                row.cells[1].text = "NA"
                        elif key == 2:
                            row.cells[1].text = rec.location if rec.location else 'NA'
                        elif key == 3:
                            row.cells[1].text = rec.client_name if rec.client_name else 'NA'
                        elif key == 6:
                            row.cells[1].text = rec.role if rec.role else 'NA'
                document.add_paragraph('')
        else:
            p.add_run('NA')
